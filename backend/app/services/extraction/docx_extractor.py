from app.core.logging import get_logger
from app.services.extraction.base import DocumentExtractor, ExtractionResult, PageResult

logger = get_logger(__name__)


class DocxExtractor(DocumentExtractor):
    def extract(self, file_data: bytes, filename: str) -> ExtractionResult:
        import io
        from docx import Document as DocxDocument

        doc = DocxDocument(io.BytesIO(file_data))

        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Extract tables
        tables = []
        for table_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                rows.append(row_data)
            if rows:
                tables.append({
                    "table_index": table_idx,
                    "rows": rows,
                    "headers": rows[0] if rows else [],
                })

        full_text = "\n".join(paragraphs)

        # Add table text
        for table in tables:
            table_text = "\n".join(
                " | ".join(str(cell) for cell in row)
                for row in table["rows"]
            )
            full_text += f"\n\n[Table {table['table_index'] + 1}]\n{table_text}"

        # DOCX doesn't have pages natively, treat as single page
        page = PageResult(
            page_number=1,
            text=full_text,
            extraction_method="python_docx",
            confidence=1.0,
            tables=tables if tables else None,
        )

        return ExtractionResult(
            pages=[page],
            total_pages=1,
            metadata={"filename": filename, "extractor": "docx", "paragraph_count": len(paragraphs)},
        )
